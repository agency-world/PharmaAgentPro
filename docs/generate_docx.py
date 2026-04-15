"""Generate TECHNICAL_GUIDE.docx from the markdown content."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    return h


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    # Add shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F5F5F5")
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def parse_table(lines):
    """Parse markdown table lines into headers and rows."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        cells = [re.sub(r"\*\*([^*]+)\*\*", r"\1", c) for c in cells]
        if i == 0:
            headers = cells
        elif i == 1:
            continue  # separator
        else:
            rows.append(cells)
    return headers, rows


def add_formatted_paragraph(doc, text, bold=False, italic=False, font_size=10):
    p = doc.add_paragraph()
    # Handle bold/italic markers in text
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(font_size)
        else:
            run = p.add_run(part)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(font_size)
    return p


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Modify default styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for i in range(1, 5):
        try:
            hs = doc.styles[f"Heading {i}"]
            hs.font.name = "Calibri"
            hs.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        except KeyError:
            pass

    # ── COVER PAGE ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PharmaAgent Pro")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Technical Guide")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Product Design, Development & Deployment")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    for _ in range(4):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Version 1.0  |  April 2026  |  Internal Technical Documentation")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # ── TABLE OF CONTENTS ──
    add_styled_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Product Overview",
        "2. System Architecture",
        "3. Agent Engineering & Configuration",
        "4. Tool System Design",
        "5. LLM Strategy & Data Retrieval",
        "6. Environment & Session Management",
        "7. Multi-Turn Conversation & Reuse",
        "8. Guardrails & Compliance",
        "9. Development Process",
        "10. Deployment Process",
        "11. Process Flow Diagrams",
        "12. FAQ & Troubleshooting",
        "13. Roadmap",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

    doc.add_page_break()

    # ── SECTION 1: PRODUCT OVERVIEW ──
    add_styled_heading(doc, "1. Product Overview", level=1)

    add_styled_heading(doc, "What is PharmaAgent Pro?", level=2)
    doc.add_paragraph(
        "PharmaAgent Pro is a multi-agent pharmaceutical intelligence platform built on "
        "Claude's Managed Agent SDK. It serves pharmaceutical researchers, clinicians, and "
        "regulatory professionals with four core capabilities."
    )

    add_table(doc,
        ["Domain", "Capability", "Key Functions"],
        [
            ["Drug Discovery", "Compound analysis", "Drug-likeness (Lipinski), ADMET profiling, CYP interactions, SAR analysis"],
            ["Clinical Trials", "Trial intelligence", "Trial status, enrollment tracking, endpoint analysis, interim results"],
            ["Regulatory Affairs", "Document generation", "IND reports, CSRs, CIOMS, DSURs, IBs -- FDA/ICH compliant"],
            ["Safety Monitoring", "Pharmacovigilance", "AE analysis, signal detection, benefit-risk assessment, SAE triage"],
        ]
    )
    doc.add_paragraph()

    add_styled_heading(doc, "Design Principles", level=2)
    principles = [
        ("Safety First", "PII redaction, prompt injection defense, compliance validation at every boundary"),
        ("Transparent Reasoning", "Extended thinking exposes Claude's reasoning chain for complex analyses"),
        ("Cost Efficiency", "Multi-model routing + prompt caching reduces LLM costs by 60-80%"),
        ("Auditability", "Every query, tool call, and response is logged in an immutable audit trail"),
        ("Domain Fidelity", "Pharma-specific tools, skill files, and guardrails ensure accurate, compliant outputs"),
    ]
    for i, (title_text, desc) in enumerate(principles, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title_text} -- ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)

    doc.add_page_break()

    # ── SECTION 2: SYSTEM ARCHITECTURE ──
    add_styled_heading(doc, "2. System Architecture", level=1)

    add_styled_heading(doc, "High-Level Architecture", level=2)
    doc.add_paragraph(
        "The system follows a layered architecture with clear separation of concerns:"
    )

    arch_layers = [
        ("User Interfaces", "CLI Chat | Web UI (port 3000) | REST API (port 8000)"),
        ("Guardrails Layer", "Input Validator | Output Filter | Compliance Checker | Rate Limiter"),
        ("PharmaOrchestrator", "Query Classifier (Haiku 4.5) routes to 4 specialized sub-agents"),
        ("Sub-Agents", "Drug Discovery | Clinical Trials | Regulatory | Safety (all Opus 4.6)"),
        ("Tool Layer", "6 pharma-specific tools with JSON Schema definitions"),
        ("Infrastructure", "Prompt Cache (5-min TTL) | Memory Store (versioned) | Audit (JSONL) | Usage Tracker"),
        ("Data Layer", "drugs_compound_library.json | clinical_trials.json | adverse_events.json | regulatory_templates.json"),
    ]
    add_table(doc, ["Layer", "Components"], arch_layers)
    doc.add_paragraph()

    add_styled_heading(doc, "Directory Structure", level=2)
    dir_struct = """ClaudeManagedAgentApp/
  src/
    main.py                 -- CLI + FastAPI entry point
    api.py                  -- REST API server (port 8000)
    agents/
      orchestrator.py       -- Query routing via Haiku classifier
      base_agent.py         -- Core LLM logic, tool loop, caching
    tools/
      drug_tools.py         -- DrugLookupTool, DrugInteractionTool
      trial_tools.py        -- TrialSearchTool, AdverseEventSearchTool, SafetySignalTool
      document_tools.py     -- DocumentGeneratorTool
    guardrails/             -- Input validation, output filtering, compliance, rate limiting
    memory/                 -- Versioned local file-based memory store
    utils/                  -- Config, logger, audit, usage tracker
    datasets/               -- 4 JSON pharma datasets
  .claude/
    skills/                 -- 4 domain SKILL.md files
    commands/               -- 4 slash command templates
  deploy/                   -- Managed agent deployment scripts + chatbot
  tests/                    -- 71+ tests across 4 files
  docs/                     -- Technical guide + marketing deck"""
    add_code_block(doc, dir_struct)

    doc.add_page_break()

    # ── SECTION 3: AGENT ENGINEERING ──
    add_styled_heading(doc, "3. Agent Engineering & Configuration", level=1)

    add_styled_heading(doc, "3.1 Multi-Model Strategy", level=2)
    doc.add_paragraph(
        "PharmaAgent Pro uses three Claude models, each selected for a specific role:"
    )
    add_table(doc,
        ["Model", "Role", "Why This Model", "Where Used"],
        [
            ["Claude Opus 4.6", "Primary reasoning", "Deepest reasoning for drug analysis, safety assessment", "All 4 sub-agents"],
            ["Claude Haiku 4.5", "Query classifier", "Fastest model for real-time routing (~100ms)", "Orchestrator classifier"],
            ["Claude Sonnet 4.6", "Managed Agent", "Balanced cost/quality for platform agent", "deploy_managed_agent.py"],
        ]
    )
    doc.add_paragraph()
    add_formatted_paragraph(doc,
        "**Cost Optimization**: By routing classification to Haiku ($0.80/M input) instead of Opus ($15/M input), "
        "classification costs are reduced by 95%."
    )

    add_styled_heading(doc, "3.2 Orchestrator Pattern", level=2)
    doc.add_paragraph(
        "The PharmaOrchestrator implements a hierarchical routing architecture. When a user query arrives, "
        "it is first sent to the Haiku 4.5 classifier which returns a domain classification with confidence score. "
        "If confidence >= 0.7 and the domain has an agent config, the query is routed to the specialized sub-agent. "
        "Otherwise, the main orchestrator agent (with all tools) handles it."
    )

    routing_flow = """User Query -> Haiku Classifier -> {"domain": "drug_discovery", "confidence": 0.92}
  |
  confidence >= 0.7 ? YES -> Drug Discovery Agent (Opus 4.6 + Extended Thinking)
                      NO  -> Main Orchestrator Agent (all tools)"""
    add_code_block(doc, routing_flow)

    add_styled_heading(doc, "3.3 Sub-Agent Configuration", level=2)
    add_table(doc,
        ["Agent", "System Prompt Focus", "Tools", "Extended Thinking"],
        [
            ["Drug Discovery", "Medicinal chemistry, ADMET, SAR analysis", "drug_lookup, drug_interaction_check", "ON"],
            ["Clinical Trials", "Trial design, biostatistics, enrollment", "trial_search, adverse_event_search, safety_signal", "ON"],
            ["Regulatory", "FDA/ICH compliance, document generation", "generate_document + search tools", "OFF"],
            ["Safety", "Signal detection, pharmacovigilance, benefit-risk", "safety_signal_analysis, adverse_event_search", "ON"],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Sub-agents are lazy-loaded: they are only instantiated when their domain is first requested, "
        "reducing memory footprint."
    )

    add_styled_heading(doc, "3.4 Extended Thinking", level=2)
    doc.add_paragraph(
        "Extended thinking is selectively enabled for domains requiring complex reasoning. "
        "When enabled, Claude is given a 10,000-token budget to reason step-by-step before generating "
        "its response. This is critical for SAR analysis, CYP metabolism chains, statistical interpretation, "
        "and benefit-risk tradeoffs. Thinking is disabled for regulatory document generation as it is formulaic."
    )

    thinking_code = """# In base_agent.py, during LLM request construction:
if use_thinking:
    request_params["thinking"] = {
        "type": "enabled",
        "budget_tokens": 10000  # Configurable via THINKING_BUDGET_TOKENS
    }"""
    add_code_block(doc, thinking_code)

    add_styled_heading(doc, "3.5 Prompt Caching Strategy", level=2)
    doc.add_paragraph(
        "The system uses Claude's prompt caching to dramatically reduce costs in multi-turn conversations. "
        "The core system prompt (~1500 tokens, stable across requests) is marked with "
        'cache_control: {"type": "ephemeral"} for a 5-minute TTL. Memory context is injected after '
        "the cached block so it stays fresh."
    )
    add_table(doc,
        ["Request", "System Prompt Cost", "Savings"],
        [
            ["First request (cache MISS)", "$18.75/M tokens (cache write)", "Baseline"],
            ["Subsequent within 5 min (cache HIT)", "$1.50/M tokens (cache read)", "90% savings"],
        ]
    )

    doc.add_paragraph()
    add_styled_heading(doc, "3.6 Skill Files", level=2)
    doc.add_paragraph(
        "Each domain has a SKILL.md file in .claude/skills/ providing deep domain expertise. "
        "These are loaded contextually when the Claude Code agent operates in the respective domain."
    )
    add_table(doc,
        ["Skill File", "Topics Covered"],
        [
            ["drug-discovery/SKILL.md", "SAR analysis, ADMET, toxicophore flagging, Lipinski Rule of Five, lead optimization"],
            ["clinical-trials/SKILL.md", "Trial design patterns, biostatistics, DSMB briefings, adaptive trials"],
            ["regulatory-docs/SKILL.md", "Document structure per FDA/ICH, regulatory pathways, submission timelines"],
            ["safety-monitoring/SKILL.md", "Signal detection, causality (WHO-UMC), Bradford Hill, benefit-risk frameworks"],
        ]
    )

    doc.add_page_break()

    # ── SECTION 4: TOOL SYSTEM ──
    add_styled_heading(doc, "4. Tool System Design", level=1)

    add_styled_heading(doc, "4.1 Tool Inventory", level=2)
    add_table(doc,
        ["Tool", "Input", "Output", "Data Source"],
        [
            ["drug_lookup", "query, field", "Compound profiles + Lipinski", "drugs_compound_library.json"],
            ["drug_interaction_check", "compound_names[]", "CYP interaction pairs + risk", "drugs_compound_library.json"],
            ["trial_search", "query, field", "Trial details + enrollment", "clinical_trials.json"],
            ["adverse_event_search", "compound, trial_id, filters", "AE records + summary stats", "adverse_events.json"],
            ["safety_signal_analysis", "compound_name", "Signal detection + SOC analysis", "adverse_events.json + clinical_trials.json"],
            ["generate_document", "doc_type, compound, fields", "Regulatory scaffold + metadata", "regulatory_templates.json"],
        ]
    )

    doc.add_paragraph()
    add_styled_heading(doc, "4.2 Agentic Tool Loop", level=2)
    doc.add_paragraph(
        "The base agent implements an agentic tool loop (max 10 iterations) that allows Claude to call "
        "tools iteratively until it has enough data to synthesize a final response:"
    )
    tool_flow = """Iteration 1: Claude calls drug_lookup(query="Nexavirin") -> compound data returned
Iteration 2: Claude calls safety_signal_analysis(compound_name="Nexavirin") -> signal data
Iteration 3: Claude calls adverse_event_search(compound_name="Nexavirin") -> AE records
Iteration 4: Claude synthesizes all data -> final text response (stop_reason: end_turn)"""
    add_code_block(doc, tool_flow)

    doc.add_paragraph(
        "Exit conditions: stop_reason != 'tool_use', or max 10 iterations reached."
    )

    add_styled_heading(doc, "4.3 Managed Agent Custom Tool Flow", level=2)
    doc.add_paragraph(
        "When deployed to platform.claude.com, custom tools execute locally on the client:"
    )
    managed_flow = """1. Agent emits 'agent.custom_tool_use' event (name, input, id)
2. Client executes tool locally against JSON datasets
3. Session goes idle with stop_reason.type == 'requires_action'
4. Client sends 'user.custom_tool_result' with custom_tool_use_id
5. Agent resumes processing with tool result"""
    add_code_block(doc, managed_flow)
    doc.add_paragraph(
        "Key point: No pharma data is uploaded to the platform. Custom tools execute entirely on the client."
    )

    doc.add_page_break()

    # ── SECTION 5: LLM STRATEGY ──
    add_styled_heading(doc, "5. LLM Strategy & Data Retrieval", level=1)

    add_styled_heading(doc, "5.1 Tool-Based Retrieval (Not RAG)", level=2)
    doc.add_paragraph(
        "PharmaAgent Pro does NOT use a traditional RAG pipeline with vector embeddings. "
        "Instead, it uses a tool-based retrieval approach where Claude decides which tool to call, "
        "tools execute deterministic searches against structured JSON datasets, and Claude synthesizes the results."
    )

    add_table(doc,
        ["Factor", "Tool-Based (Our Approach)", "RAG with Embeddings"],
        [
            ["Data size", "4 focused JSON files (~50KB)", "Better for millions of docs"],
            ["Precision", "Exact field matching", "Semantic similarity (may miss)"],
            ["Transparency", "Claude names the tool it calls", "Retrieval is opaque"],
            ["Structured data", "Native JSON filtering", "Requires flattening"],
            ["Compliance", "Tool calls fully auditable", "Harder to audit"],
            ["Cost", "No embedding costs", "Requires embedding API"],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "When RAG would be appropriate: If the dataset grows to thousands of documents "
        "(e.g., full FDA label database, literature corpus), a vector store with semantic search "
        "would be added as an additional tool."
    )

    add_styled_heading(doc, "5.2 Memory Store", level=2)
    doc.add_paragraph(
        "The memory store (src/memory/memory_store.py) provides persistent context across sessions. "
        "It uses a local file-based store with automatic versioning. Before each LLM call, "
        "get_context_for_agent() searches for memories tagged with the active agent and injects them "
        "into the system prompt after the cached block."
    )
    mem_struct = """store_root/
  index.json                 -- Master index
  memories/<hash>.json       -- Current state
  versions/<hash>/v1.json    -- Immutable version history"""
    add_code_block(doc, mem_struct)

    add_styled_heading(doc, "5.3 Prompt Caching (API-Level)", level=2)
    doc.add_paragraph(
        "Prompt caching is a Claude API feature, not a local cache layer. The system prompt is written to "
        "cache on the first request ($18.75/M for Opus) and subsequent requests within 5 minutes read from "
        "cache ($1.50/M -- a 90% savings). The usage tracker records cache_creation_tokens and cache_read_tokens "
        "separately for cost analysis."
    )

    doc.add_page_break()

    # ── SECTION 6: ENVIRONMENT & SESSION ──
    add_styled_heading(doc, "6. Environment & Session Management", level=1)

    add_styled_heading(doc, "6.1 Managed Agent Environment", level=2)
    doc.add_paragraph(
        "The deployment creates a cloud environment on platform.claude.com with pip packages "
        "(pandas, jinja2, rich) and unrestricted networking. This environment provides an isolated "
        "runtime where the managed agent can execute built-in tools."
    )

    add_styled_heading(doc, "6.2 Session Lifecycle", level=2)
    session_flow = """Create Session -> status_idle (end_turn)
  -> User sends message -> status_running
    -> agent.thinking -> agent.message -> agent.custom_tool_use
    -> status_idle (requires_action) -> Client sends tool result
    -> status_running (resumes) -> agent.message
    -> status_idle (end_turn)
  -> Ready for next message"""
    add_code_block(doc, session_flow)

    add_styled_heading(doc, "6.3 Local Session Management", level=2)
    doc.add_paragraph(
        "The local API server maintains per-session orchestrator instances in memory. "
        "Each session has independent conversation history, sub-agent instances (lazy-loaded), "
        "memory context, and usage tracking."
    )

    doc.add_page_break()

    # ── SECTION 7: MULTI-TURN ──
    add_styled_heading(doc, "7. Multi-Turn Conversation & Reuse", level=1)

    add_styled_heading(doc, "7.1 Conversation Context Retention", level=2)
    doc.add_paragraph(
        "Multi-turn conversation is fully enabled. The full message history (self.messages) is maintained "
        "and sent with every API call. Claude naturally resolves references like 'its', 'that compound' from prior turns."
    )

    multi_turn = """Turn 1: "Tell me about Nexavirin"
  -> Agent calls drug_lookup, returns compound profile
  -> Messages: [user_1, assistant_1]

Turn 2: "What are its CYP interactions?"
  -> 'its' resolves to Nexavirin from context
  -> Agent calls drug_interaction_check
  -> Messages: [user_1, assistant_1, user_2, assistant_2]

Turn 3: "Compare that with Oncolytin-B"
  -> 'that' resolves to CYP interaction profile
  -> Full context from prior turns available"""
    add_code_block(doc, multi_turn)

    add_styled_heading(doc, "7.2 How Similar Questions Benefit", level=2)
    doc.add_paragraph("When a user asks follow-up or similar questions, the system benefits in three ways:")
    benefits = [
        ("Conversation History", "Prior tool results are already in the message history. Claude can reference them without re-calling tools."),
        ("Prompt Cache Hits", "System prompt cached for 5 minutes. Follow-up questions reuse the cache, reducing costs by ~90% on that portion."),
        ("Memory Store", "If the agent saves key findings to memory during a session, subsequent sessions retrieve that context automatically."),
    ]
    for title_text, desc in benefits:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)

    doc.add_page_break()

    # ── SECTION 8: GUARDRAILS ──
    add_styled_heading(doc, "8. Guardrails & Compliance", level=1)

    add_styled_heading(doc, "8.1 Guardrail Pipeline", level=2)
    doc.add_paragraph(
        "Every request passes through a four-stage guardrail pipeline:"
    )
    pipeline = """INPUT -> [1] Rate Limiter (30 req/min per user)
      -> [2] Input Validator (PII + injection + prohibited topics)
      -> [3] LLM Processing (Claude generates response)
      -> [4] Output Filter (PII redaction + medical advice warnings + compliance footer)
      -> OUTPUT"""
    add_code_block(doc, pipeline)

    add_styled_heading(doc, "8.2 PII Handling", level=2)
    add_table(doc,
        ["PII Type", "Pattern", "Replacement"],
        [
            ["SSN", "XXX-XX-XXXX", "[REDACTED_SSN]"],
            ["Email", "Standard email regex", "[REDACTED_EMAIL]"],
            ["Phone", "XXX-XXX-XXXX", "[REDACTED_PHONE]"],
            ["Medical Record #", "MRN[-:]?\\s*\\d+", "[REDACTED_MRN]"],
            ["Date of Birth", "DOB[-:]?\\s*date", "[REDACTED_DOB]"],
            ["Credit Card", "XXXX-XXXX-XXXX-XXXX", "[REDACTED_CC]"],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Guarantee: PII is redacted before being sent to the Claude API and again after response generation."
    )

    add_styled_heading(doc, "8.3 Compliance Validation", level=2)
    add_table(doc,
        ["Document Type", "Standard", "Validates"],
        [
            ["IND Annual Report", "21 CFR 312.33", "ind_number, sponsor, 8 required sections"],
            ["Clinical Study Report", "ICH E3", "study_id, phase, 10 required sections"],
            ["CIOMS Safety Report", "ICH E2A", "ae_id, event_term, 4 required sections"],
            ["DSUR", "ICH E2F", "development_phase, DIBD, 8 required sections"],
            ["Investigator's Brochure", "ICH E6(R2)", "version, date, 5 required sections"],
        ]
    )

    doc.add_page_break()

    # ── SECTION 9: DEVELOPMENT ──
    add_styled_heading(doc, "9. Development Process", level=1)

    add_styled_heading(doc, "9.1 Setup", level=2)
    setup_code = """git clone <repo-url> && cd ClaudeManagedAgentApp
python -m venv .venv && source .venv/bin/activate
pip install -e ".[dev]"
cp .env.example .env   # Set ANTHROPIC_API_KEY"""
    add_code_block(doc, setup_code)

    add_styled_heading(doc, "9.2 Testing", level=2)
    doc.add_paragraph("The test suite contains 71+ tests across 4 files (no API key needed for unit tests):")
    add_table(doc,
        ["Test File", "Count", "Coverage"],
        [
            ["test_tools.py", "30+", "All 6 tools: lookup, interaction, search, AE, signal, document"],
            ["test_guardrails.py", "20+", "Input validation, output filtering, compliance, rate limiting"],
            ["test_memory.py", "10+", "Write/read, versioning, search, tags, agent context"],
            ["test_utils.py", "15+", "Config, audit logger, usage tracker, cost calculation"],
        ]
    )

    doc.add_paragraph()
    add_styled_heading(doc, "9.3 Local Development Modes", level=2)
    add_table(doc,
        ["Mode", "Command", "Description"],
        [
            ["Interactive CLI", "python -m src.main", "Multi-turn chat with slash commands"],
            ["API Server", "python -m src.main --mode api", "FastAPI on port 8000 with /docs"],
            ["Managed Agent (Terminal)", "python deploy/chatbot.py", "Connects to platform.claude.com"],
            ["Managed Agent (Web)", "python deploy/chatbot_server.py", "Web UI on port 3000"],
        ]
    )

    doc.add_page_break()

    # ── SECTION 10: DEPLOYMENT ──
    add_styled_heading(doc, "10. Deployment Process", level=1)

    add_styled_heading(doc, "10.1 Deployment Sequence", level=2)
    doc.add_paragraph("The deployment to platform.claude.com follows a 4-step process:")

    deploy_steps = [
        ("Step 1: Create Agent", "POST /v1/agents -- Name, model (Sonnet 4.6), system prompt, built-in toolset + 6 custom tools with JSON Schema"),
        ("Step 2: Create Environment", "POST /v1/environments -- Cloud type, pip packages (pandas, jinja2, rich), unrestricted networking"),
        ("Step 3: Create Memory Store", "POST /v1/memory_stores -- Seed with pharma context. Currently in research preview; gracefully degraded."),
        ("Step 4: Create Session", "POST /v1/sessions -- Links agent_id + environment_id. Optionally attaches memory_store."),
    ]
    for step_title, step_desc in deploy_steps:
        p = doc.add_paragraph()
        run = p.add_run(f"{step_title}: ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(step_desc)
        run2.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph(
        "State is persisted to deploy_state.json with all resource IDs. "
        "The deployment script is idempotent -- it skips existing resources."
    )

    add_styled_heading(doc, "10.2 Running the Deployment", level=2)
    deploy_code = """export ANTHROPIC_API_KEY=sk-ant-...
python deploy/deploy_managed_agent.py    # Deploy (idempotent)
python deploy/chatbot.py                 # Terminal chatbot
python deploy/chatbot_server.py          # Web UI at http://localhost:3000"""
    add_code_block(doc, deploy_code)

    add_styled_heading(doc, "10.3 Architecture (Local Client + Platform)", level=2)
    doc.add_paragraph(
        "Custom tools execute locally on the client. The managed agent sends tool call requests; "
        "the client executes them against local datasets and returns results. "
        "No pharma data is uploaded to the platform."
    )

    doc.add_page_break()

    # ── SECTION 11: PROCESS FLOWS ──
    add_styled_heading(doc, "11. Process Flow Diagrams", level=1)

    add_styled_heading(doc, "11.1 End-to-End Request Flow", level=2)
    e2e_flow = """User: "Run safety signal analysis on Oncolytin-B"
  |
  [Rate Limiter] PASS (23 remaining)
  [Input Validator] PII: none, Injection: clean, Topics: clean
  |
  [Orchestrator] -> [Haiku Classifier]
    Classification: {"domain": "safety", "confidence": 0.95}
    Route: Safety Agent (Opus 4.6 + Extended Thinking)
  |
  [Safety Agent] -> API Call #1
    Claude thinks: "I need to analyze Oncolytin-B safety..."
    Claude calls: safety_signal_analysis(compound_name="Oncolytin-B")
  |
  [Tool Executor] Loads AE + trial data, calculates signals
    Returns: ALERT (ILD signal, 2 SAEs, 2 Grade 3+)
  |
  [Safety Agent] -> API Call #2
    Claude calls: adverse_event_search(serious_only=true)
    Returns: 2 SAE records with narratives
  |
  [Safety Agent] -> API Call #3
    Claude synthesizes all data -> final response (stop_reason: end_turn)
  |
  [Output Filter] PII: 0 redactions, Compliance footer added
  [Audit Logger] Logs: query, 2 tool_use, response events
  [Usage Tracker] 3 API calls, ~5000 input tokens, ~$0.19 estimated
  |
  Response displayed to user"""
    add_code_block(doc, e2e_flow)

    add_styled_heading(doc, "11.2 Managed Agent Custom Tool Flow", level=2)
    custom_flow = """Client                              Platform
  |                                      |
  | -- user.message event ------------> |
  |                                      | Agent processes
  | <-- agent.thinking ---------------- |
  | <-- agent.custom_tool_use ---------- | (drug_lookup, id=evt_abc)
  |                                      |
  Execute tool locally                   | Session goes idle
  | <-- session.status_idle ------------ | stop_reason: requires_action
  |     event_ids: ["evt_abc"]           |
  |                                      |
  | -- user.custom_tool_result -------> |
  |    custom_tool_use_id: "evt_abc"     |
  |    content: [{text: "..."}]          |
  |                                      | Agent resumes
  | <-- agent.message ------------------ | Final response
  | <-- session.status_idle ------------ | stop_reason: end_turn"""
    add_code_block(doc, custom_flow)

    doc.add_page_break()

    # ── SECTION 12: FAQ ──
    add_styled_heading(doc, "12. FAQ & Troubleshooting", level=1)

    add_styled_heading(doc, "Frequently Asked Questions", level=2)
    faqs = [
        ("Does PharmaAgent Pro use RAG?",
         "No. It uses tool-based retrieval where Claude decides which tool to call, and tools execute "
         "deterministic searches against structured JSON datasets. This provides exact matching, full "
         "auditability, and zero embedding costs."),
        ("Is patient data sent to Claude?",
         "No. PII is detected and redacted at the input boundary before any API call. Output is also filtered. "
         "The audit trail logs all redactions."),
        ("How does multi-turn context work?",
         "Full conversation history is maintained in self.messages and sent with every API call. Claude "
         "naturally resolves references like 'its', 'that compound' from prior turns."),
        ("Why Sonnet for the managed agent instead of Opus?",
         "The managed agent on platform.claude.com handles general orchestration. Complex domain reasoning "
         "happens in the local sub-agents (Opus). Sonnet provides balanced quality/cost."),
        ("How are costs controlled?",
         "Three mechanisms: (1) Haiku for classification (95% cheaper), (2) prompt caching (90% savings), "
         "(3) rate limiting (30 req/min per user)."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f"Q: {q}")
        run.bold = True
        run.font.size = Pt(10)
        p2 = doc.add_paragraph(f"A: {a}")
        p2.paragraph_format.space_after = Pt(8)

    add_styled_heading(doc, "Common Errors", level=2)
    add_table(doc,
        ["Error", "Cause", "Fix"],
        [
            ["ANTHROPIC_API_KEY not set", "Missing .env or env var", "Set key in .env file"],
            ["tool_use_id: Extra inputs not permitted", "Wrong field name", "Use custom_tool_use_id"],
            ["Rate limit exceeded", "Too many requests", "Wait or increase RATE_LIMIT_PER_MINUTE"],
            ["Session terminated", "Session expired", "Create new session via /api/new-session"],
            ["Memory Store API not available", "Research preview", "Expected -- uses local memory fallback"],
        ]
    )

    doc.add_page_break()

    # ── SECTION 13: ROADMAP ──
    add_styled_heading(doc, "13. Roadmap", level=1)

    add_styled_heading(doc, "Phase 1: Foundation Hardening (Current + Next 2 Months)", level=2)
    phase1 = [
        "Vector Search Integration -- Embedding-based retrieval for large document corpora",
        "Streaming Responses -- Token-by-token streaming in CLI and Web UI",
        "Session Persistence -- Conversation history saved to disk for cross-restart recovery",
        "Managed Agent Memory Store -- Adopt platform API when it exits research preview",
        "Error Recovery -- Retry logic with exponential backoff for transient API errors",
    ]
    for item in phase1:
        doc.add_paragraph(item, style="List Bullet")

    add_styled_heading(doc, "Phase 2: Multi-Agent Orchestration (Months 3-5)", level=2)
    phase2 = [
        "Parallel Agent Execution -- Run drug_discovery + safety agents concurrently",
        "Agent-to-Agent Communication -- Safety agent can request data from clinical trials agent",
        "Supervisor Agent -- Meta-agent reviews sub-agent outputs for consistency",
        "Dynamic Tool Discovery -- Agents request tools from other agents at runtime",
        "Conflict Resolution -- Escalate disagreements to supervisor agent",
    ]
    for item in phase2:
        doc.add_paragraph(item, style="List Bullet")

    add_styled_heading(doc, "Phase 3: Enterprise Features (Months 6-9)", level=2)
    phase3 = [
        "Role-Based Access Control -- Different tool permissions per role",
        "Real-Time Data Connectors -- FDA FAERS, ClinicalTrials.gov, PubMed APIs",
        "Document Versioning & Approval Workflows -- Revision tracking with approval chains",
        "Multi-Tenant Support -- Isolated memory, audit, and sessions per org",
        "Scheduled Monitoring -- Automated daily safety scans with alert notifications",
    ]
    for item in phase3:
        doc.add_paragraph(item, style="List Bullet")

    add_styled_heading(doc, "Phase 4: Advanced Intelligence (Months 9-12)", level=2)
    phase4 = [
        "Literature-Augmented Analysis -- RAG over biomedical literature",
        "Predictive Safety Modeling -- Predict signals from historical AE data",
        "Automated Regulatory Submission Prep -- End-to-end IND/NDA package assembly",
        "Multi-Modal Analysis -- Molecular structure images, clinical charts, lab reports",
        "Federated Learning -- Cross-org safety signal sharing without raw data exposure",
    ]
    for item in phase4:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    add_styled_heading(doc, "Deployment Milestones", level=2)
    add_table(doc,
        ["Milestone", "Target", "Description"],
        [
            ["MVP", "Done", "Local + managed agent, 6 tools, 4 agents"],
            ["CI/CD Pipeline", "Month 1", "GitHub Actions automated testing + deploy"],
            ["Staging Environment", "Month 2", "Pre-production validation"],
            ["Production", "Month 3", "Docker/K8s with monitoring + autoscaling"],
            ["SOC 2 Compliance", "Month 6", "Audit hardening, encryption, access logging"],
            ["GxP Validation", "Month 9", "21 CFR Part 11 for regulated use"],
        ]
    )

    # ── FOOTER ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("-- End of Technical Guide --")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.italic = True

    return doc


if __name__ == "__main__":
    print("Generating TECHNICAL_GUIDE.docx...")
    doc = build_document()
    output = Path(__file__).parent / "PharmaAgent_Pro_Technical_Guide.docx"
    doc.save(str(output))
    print(f"Saved to {output}")
